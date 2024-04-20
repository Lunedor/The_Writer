import os
import webview
import threading
from flask import Flask, render_template, request, jsonify, send_from_directory, url_for
import json
import re
from flask_weasyprint import HTML
from docx import Document
from docx.shared import Inches
import requests
from io import BytesIO
from openai import OpenAI
import google.generativeai as genai
from werkzeug.utils import secure_filename
import datetime

app = Flask(__name__)

home_directory = os.getenv('APPDATA')

UPLOAD_FOLDER = os.path.join(home_directory, ".the-writer")
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
FOLDERS_FILE = UPLOAD_FOLDER + '/folders.json'
PREFERENCES_FILE_PATH = UPLOAD_FOLDER + '/preferences.json'
API_KEY_FILE = UPLOAD_FOLDER + '/api_key.json'
THEME_FILE_PATH = UPLOAD_FOLDER + '/colors.json'
ALLOWED_EXTENSIONS = {'txt', 'html', 'docx', 'csv', 'md', 'pdf'}
# Configuration variables
ALLOWED_IMAGE_TYPES = ['image/png', 'image/jpeg', 'image/jpg']
IMAGE_MAX_SIZE = 1024 * 1024 * 2  # 2 MB

home_directory = os.getenv('APPDATA')
UPLOAD_IMAGE_FOLDER = os.path.join(home_directory, ".the-writer", "images")
os.makedirs(UPLOAD_IMAGE_FOLDER, exist_ok=True)

@app.route('/imageUpload', methods=['POST'])
def upload_image():
    if 'image' not in request.files:
        return jsonify({'error': 'noFileGiven'}), 400

    file = request.files['image']
    if file.filename == '':
        return jsonify({'error': 'noFileGiven'}), 400

    if file and file.content_type in ALLOWED_IMAGE_TYPES:
        if file.content_length > IMAGE_MAX_SIZE:
            return jsonify({'error': 'fileTooLarge'}), 413

        # Adding a timestamp to filename
        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
        filename = secure_filename(f"{timestamp}_{file.filename}")
        file_path = os.path.join(UPLOAD_IMAGE_FOLDER, filename)
        file.save(file_path)
        print(f"Image saved to: {file_path}")

        # Generate the URL for the uploaded file
        protocol = 'http://' if request.environ.get('HTTPS') == 'off' else 'https://'
        host = request.host
        image_url = f"images/{filename}"
        print(image_url)
        response_data = {
            'data': {
                'filePath': image_url
            }
        }

        return jsonify(response_data), 200

@app.route('/images/<filename>')
def serve_image(filename):
    return send_from_directory(UPLOAD_IMAGE_FOLDER, filename)
    
@app.route('/save_api_key', methods=['POST'])
def save_api_key():
    data = request.get_json()
    model_base = data.get('model_type')
    model_type = "openai" if model_base in ["gpt-3.5-turbo", "gpt-4-turbo"] else "gemini"
    api_key = data.get('api_key')

    if not model_type or not api_key:
        return jsonify({"error": "Missing model type or API key"}), 400

    try:
        # Check if the file exists and is not empty
        if not os.path.exists(API_KEY_FILE) or os.stat(API_KEY_FILE).st_size == 0:
            with open(API_KEY_FILE, 'w') as f:
                json.dump({}, f)  # Create a new JSON file if not present

        with open(API_KEY_FILE, 'r+') as f:
            keys = json.load(f)
            keys[model_type] = api_key  # Add or update the model type key
            f.seek(0)  # Reset file pointer to the beginning of the file
            json.dump(keys, f)
            f.truncate()  # Remove any leftover data

        return jsonify({"message": "API Key saved successfully!"}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/get_api_key', methods=['GET'])
def get_api_key():
    model_type = request.args.get('model_type')  # Query parameter to specify model type
    try:
        with open(API_KEY_FILE, 'r') as f:
            keys = json.load(f)
            api_key = keys.get(model_type, '')
        return jsonify({"api_key": api_key}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def load_api_key(model_type):
    """ Load the API key for the specified model type from the stored JSON file """
    try:
        with open(API_KEY_FILE, 'r') as file:
            data = json.load(file)
            return data.get(model_type, '')  # Return the API key based on model type
    except FileNotFoundError:
        return ""  # Return an empty string if the file does not exist

def call_gemini_api(prompt, api_key):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-pro')  # Modify as needed
    response = model.generate_content(prompt)
    return response.text

def call_openai_api(prompt, api_key, model_version):       
    data = request.json # Load the API key dynamically
    api_key = load_api_key('openai')  # Load the API key dynamically
    client = OpenAI(api_key=api_key)  # Initialize the OpenAI client with the loaded API key

    messages = [
        {"role": "system", "content": "You are a writing assistant. AYou always write same language as the given text. Your main task is to provide a new text that provides continuity to the text. Unless specifically requested, these texts are not the end or summary of the story, but a chapter within the book. You are an expert at capturing the unique characteristics of an author from a text and maintaining the writer's style."},
        {"role": "user", "content": data['prompt']}
    ]

    try:
        chat_completion = client.chat.completions.create(
            messages=messages,
            model=data['model_version'],
            max_tokens=4096
        )
        response_content = chat_completion.choices[0].message.content
        return response_content
    except Exception as e:
        print("Error:", str(e))  # Print any errors
        return jsonify({"error": str(e)}), 500
        
@app.route('/get_ai_response', methods=['POST'])
def get_ai_response():
    data = request.json
    model_version = data['model_version']
    prompt = data['prompt']
    
    try:
        if "gemini" in model_version:
            api_key = load_api_key('gemini')  # Specify 'gemini' for Gemini API key
            response = call_gemini_api(prompt, api_key)
        else:
            api_key = load_api_key('openai')  # Specify 'openai' for OpenAI API key
            response = call_openai_api(prompt, api_key, model_version)
    
        return jsonify({"response": response}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/save-theme', methods=['POST'])
def save_theme():
    data = request.json
    with open(THEME_FILE_PATH, 'w') as f:
        json.dump(data, f)
    return jsonify({"message": "Theme saved successfully"})

@app.route('/load-theme', methods=['GET'])
def load_theme():
    if os.path.exists(THEME_FILE_PATH):
        with open(THEME_FILE_PATH, 'r') as f:
            theme_data = json.load(f)
    else:
        # Default theme values from your initial themeData structure
        theme_data = {
            "default": {
                "light": {
                    "--bg-color-light": "#ffffe8",
                    "--text-color-light": "#222222",
                    "--hover-color-light": "#cccccc",
                    "--editor-bg-color-light": "#ffffee",
                    "--editor-text-color-light": "#000000",
                    "--editor-text-color-dimming-light": "#E9E9E9",
                },
                "dark": {
                    "--bg-color-dark": "#111111",
                    "--text-color-dark": "#ffffff",
                    "--hover-color-dark": "#444444",
                    "--editor-bg-color-dark": "#222222",
                    "--editor-text-color-dark": "#ffffff",
                    "--editor-text-color-dimming-dark": "#222222",
                }
            },
            "custom": []  # Initialize custom as an empty array
        }
    return jsonify(theme_data)

@app.route('/delete-theme', methods=['POST'])
def delete_theme():
    data = request.json
    theme_name_to_delete = data.get('themeName')

    if not theme_name_to_delete or theme_name_to_delete == 'default':
        return jsonify({'error': 'Cannot delete the default theme or no theme name provided.'}), 400

    if os.path.exists(THEME_FILE_PATH):
        with open(THEME_FILE_PATH, 'r+') as f:
            themes_data = json.load(f)
            # Assuming your custom themes are stored under 'custom', not 'themes'
            custom_themes = themes_data.get('custom', [])
            
            # Find the theme by name
            theme_to_delete = next((theme for theme in custom_themes if theme['name'] == theme_name_to_delete), None)
            
            if theme_to_delete:
                # Remove the theme from the list
                themes_data['custom'] = [theme for theme in custom_themes if theme['name'] != theme_name_to_delete]
                
                # Rewriting the updated themes back to the file
                f.seek(0)  # Move to the start of the file
                json.dump(themes_data, f)
                f.truncate()  # Truncate the file to the new length
                return jsonify({'message': f'Theme {theme_name_to_delete} deleted successfully.'})
            else:
                return jsonify({'error': 'Theme not found.'}), 404
    return jsonify({'error': 'Theme data file not found.'}), 500

@app.route('/preferences', methods=['GET'])
def get_preferences():
    try:
        with open(PREFERENCES_FILE_PATH, 'r') as f:
            preferences = json.load(f)
    except FileNotFoundError:
        preferences = {}  # Default preferences if file does not exist
    return jsonify(preferences)

@app.route('/preferences', methods=['POST'])
def save_preferences():
    preferences = request.json
    with open(PREFERENCES_FILE_PATH, 'w') as f:
        json.dump(preferences, f, indent=4)
    return jsonify({'status': 'success'})

@app.route('/get-markdown-headers', methods=['POST'])
def get_markdown_headers():
    data = request.json
    filePath = data.get('filePath')
    # If you're expecting a path, make sure to read the file's contents
    try:
        with open(filePath, 'r', encoding='utf-8') as file:
            content = file.read()
        headers = parse_markdown_headers(content)  # Your function to extract headers
        return jsonify(headers)
    except Exception as e:
        print(f"An error occurred: {str(e)}")
        return jsonify({'error': 'Failed to process the request'}), 500

def parse_markdown_headers(markdown_content):
    headers = []
    current_level_nodes = {0: headers}  # Root level
    line_number = 0

    for line in markdown_content.split('\n'):
        line_number += 1  # Increment line number for each line processed
        match = re.match(r'^(#+)\s+(.*)', line)
        if match:
            level = len(match.group(1))
            header_text = match.group(2).strip()
            new_header = {
                "text": header_text, 
                "level": level, 
                "children": [],
                "line": line_number  # Include the line number in the header's data
            }

            parent_level = level - 1
            if parent_level in current_level_nodes:
                current_level_nodes[parent_level].append(new_header)
            else:
                print(f"Unexpected header level: {level} for header: {header_text}")

            current_level_nodes[level] = new_header['children']

    return headers

@app.route('/select-folder', methods=['GET'])
def select_folder():
    """Open a dialog to select a folder and return the selected path."""
    window = webview.windows[0]
    folder_path = window.create_file_dialog(webview.FOLDER_DIALOG)
    if folder_path:
        return jsonify({"folderPath": folder_path[0]}), 200
    else:
        return jsonify({"message": "No folder selected"}), 400

def read_folders():
    """Read the folder structure from the JSON file."""
    try:
        with open(FOLDERS_FILE, 'r') as file:
            return json.load(file)
    except FileNotFoundError:
        return []

def write_folders(folders):
    """Write the folder structure to the JSON file."""
    with open(FOLDERS_FILE, 'w') as file:
        json.dump(folders, file, indent=4)

@app.route('/add-folder', methods=['POST'])
def add_folder():
    data = request.json
    new_folder = {'name': data.get('folderName'), 'path': data.get('folderPath')}

    folders = read_folders()
    if new_folder['path'] in [folder['path'] for folder in folders]:
        return jsonify({'message': 'This folder is already added.'}), 400

    folders.append(new_folder)
    write_folders(folders)

    return jsonify({'message': 'Folder added successfully.'}), 200

@app.route('/files-with-structure')
def get_files_with_structure():
    # Directly read from folders.json and return its content
    return jsonify(read_folders())

@app.route('/folder-files', methods=['GET'])
def get_folder_files():
    folder_path = request.args.get('folderPath', '')
    
    # Retrieve stored folders
    folders = read_folders()
    folder = next((f for f in folders if f['path'] == folder_path), None)

    if not folder:
        return jsonify({'message': 'Folder not found'}), 404

    try:
        # List only .md files if the folder exists in the stored data
        files = [f for f in os.listdir(folder_path) 
                 if os.path.isfile(os.path.join(folder_path, f)) and f.endswith('.md')]
        return jsonify({'files': files})
    except FileNotFoundError:
        return jsonify({'message': 'Folder not found in filesystem'}), 404

@app.route('/remove-folder', methods=['POST'])
def remove_folder():
    data = request.json
    folder_path = data.get('folderPath')
    folders = read_folders()
    
    # Filter out the folder to remove
    updated_folders = [folder for folder in folders if folder['path'] != folder_path]
    write_folders(updated_folders)  # Save the updated list back to folders.json
    
    return jsonify({'message': 'Folder removed successfully.'}), 200

@app.route('/')
def index():
    return render_template('index.html')

def start_server():
    app.run(host='127.0.0.1', port=5000)
    
# Python: Adjust save_file to handle directory paths
@app.route('/save', methods=['POST'])
def save_file():
    data = request.json
    filename = data['filename']
    content = data['content']
    # Save the content to a file in Markdown format
    with open(f'{filename}.md', 'w', encoding='utf-8') as file:
        file.write(content)
    return jsonify({"message": "File saved successfully!"})


@app.route('/load', methods=['POST'])
def load_file():
    data = request.json
    filename = data['filename']
    
    # Append '.md' if it's not part of the filename already
    if not filename.lower().endswith('.md'):
        filename += '.md'

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    if not os.path.exists(file_path):
        return jsonify({'message': 'File not found.'}), 404

    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        return jsonify({'content': content}), 200
    except IOError as e:
        return jsonify({'error': str(e)}), 500

@app.route('/save-as', methods=['POST'])
def save_as_file():
    data = request.json
    file_name = data['filename']
    file_format = data['fileFormat']
    content = data['content']
    
    file_path = webview.windows[0].create_file_dialog(
        webview.SAVE_DIALOG, 
        directory=os.path.expanduser('~'), 
        save_filename=f"{file_name}.{file_format}")
    
    try:
        if file_format == 'docx':
            document = Document()

            # Split content to handle elements one by one
            content_items = re.split('(<img.*?>)', content)  # Split by image tags to maintain order
            
            for item in content_items:
                if '<img' in item:  # If the item is an image
                    image_url_match = re.search(r'src="(.*?)"', item)
                    style_match = re.search(r'style="(.*?)"', item)
                    if image_url_match:
                        url = image_url_match.group(1)
                        width_inch = None
                        if style_match:
                            width_style = style_match.group(1)
                            width_px_match = re.search(r'width:\s*(\d+)px', width_style)
                            if width_px_match:
                                width_px = int(width_px_match.group(1))
                                width_inch = width_px / 96  # Convert pixels to inches
                        response = requests.get(url)
                        image_io = BytesIO(response.content)
                        document.add_picture(image_io, width=Inches(width_inch) if width_inch else None)
                else:  # The item is text
                    text = re.sub('<.*?>', '', item).strip()  # Remove HTML tags from text
                    if text:  # Avoid adding empty paragraphs
                        document.add_paragraph(text)

            # Save the document
            document.save(file_path)
        
        elif file_format == 'pdf':
            # Convert HTML to PDF using Flask-WeasyPrint
            HTML(string=content).write_pdf(file_path)
        
        elif file_format in ['txt', 'md']:
            # Save as plain text or Markdown
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(content)
        
        return jsonify({'message': 'File saved successfully.'}), 200

    except Exception as e:
        print(f"Error saving file: {e}")
        return jsonify({'message': 'Error saving file'}), 500
    

if __name__ == '__main__':
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
        
    # Start the Flask server in a separate thread
    flask_thread = threading.Thread(target=start_server, daemon=True)
    flask_thread.start()
    
    
    # Create and start the PyWebView window
    webview.create_window("The Writer", "http://127.0.0.1:5000", width=1200, height=800, confirm_close=True)
    webview.start()